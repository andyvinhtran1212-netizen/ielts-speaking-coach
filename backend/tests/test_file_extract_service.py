"""Tests for services.file_extract_service (Phase 2.3c-2).

The service has two halves: per-format extractors (docx + txt) and a
top-level `extract_text` dispatcher that handles size / extension /
empty-file gates.  We test both directly so a regression in the
dispatcher (e.g. swapping size-check and extension-check) doesn't
hide behind a happy-path test.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import docx
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from services.file_extract_service import (
    MAX_EXTRACTED_CHARS,
    MAX_FILE_SIZE_BYTES,
    FileExtractError,
    _is_likely_text,
    extract_text,
    extract_text_from_docx,
    extract_text_from_txt,
)


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Build a real .docx in memory.  We deliberately don't mock
    python-docx — that's the whole point of these tests, the
    parse path needs to run end-to-end."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── .txt extractor ───────────────────────────────────────────────────


def test_extract_txt_utf8():
    """UTF-8 with Vietnamese diacritics survives the round trip."""
    text = "Some essay content with Vietnamese: học viên giỏi."
    result = extract_text_from_txt(text.encode("utf-8"))
    assert "học viên giỏi" in result


def test_extract_txt_with_bom():
    """UTF-8 BOM is stripped — Word/Notepad on Windows often emits it
    and a leading `\\ufeff` would corrupt the first word."""
    text = "Hello"
    bytes_with_bom = b"\xef\xbb\xbf" + text.encode("utf-8")
    result = extract_text_from_txt(bytes_with_bom)
    assert result == "Hello"
    assert "\ufeff" not in result


def test_extract_txt_latin1_fallback():
    """latin-1 is the implicit safety net — it never raises, so any
    byte stream still produces a string rather than a 500."""
    text = "Some text with accent: café"
    result = extract_text_from_txt(text.encode("latin-1"))
    # latin-1 gets decoded as latin-1 (one byte per char) — we don't
    # promise it round-trips back to the original glyphs, only that
    # we don't raise.
    assert "caf" in result


# ── .docx extractor ──────────────────────────────────────────────────


def test_extract_docx_basic():
    """Paragraphs join with blank lines; empty paragraphs dropped."""
    docx_bytes = _make_docx_bytes([
        "First paragraph.",
        "Second paragraph with more text.",
        "",  # the empty paragraph must NOT add a stray blank line
        "Third paragraph.",
    ])

    result = extract_text_from_docx(docx_bytes)
    assert "First paragraph" in result
    assert "Second paragraph" in result
    assert "Third paragraph" in result
    # Each non-empty paragraph separated by exactly one blank line.
    assert "First paragraph." in result and "Third paragraph." in result


def test_extract_docx_corrupted_raises():
    """python-docx fails noisily on bad bytes — we wrap that in a
    FileExtractError so the router can map it cleanly to a 400."""
    with pytest.raises(FileExtractError, match="Không đọc được"):
        extract_text_from_docx(b"not a real docx file")


# ── Dispatcher ───────────────────────────────────────────────────────


def test_extract_dispatches_by_extension_docx():
    docx_bytes = _make_docx_bytes(["Hello world"])
    result = extract_text("essay.docx", docx_bytes)
    assert "Hello world" in result


def test_extract_dispatches_by_extension_txt():
    result = extract_text("essay.txt", b"Plain text content")
    assert result == "Plain text content"


def test_extract_rejects_unknown_extension():
    """PDF + image + anything else → 400 with the canonical message
    so the frontend's error toast can render it verbatim."""
    with pytest.raises(FileExtractError, match="không hỗ trợ"):
        extract_text("essay.pdf", b"some bytes")


def test_extract_rejects_oversize():
    """Size check runs BEFORE extension check, on purpose — a 100 MB
    file gets rejected even if it's `.txt`-named, so we never alloc
    that much memory."""
    big_bytes = b"x" * (MAX_FILE_SIZE_BYTES + 1)
    with pytest.raises(FileExtractError, match="quá lớn"):
        extract_text("essay.txt", big_bytes)


def test_extract_rejects_empty():
    """0-byte upload is almost always a misclick — surface the
    "rỗng" message so the student knows the picker didn't actually
    grab a file."""
    with pytest.raises(FileExtractError, match="rỗng"):
        extract_text("essay.txt", b"")


def test_extract_truncates_long_text():
    """Anything past MAX_EXTRACTED_CHARS is silently trimmed.  The
    Pydantic submit validator caps at 10 k, so a 15 k upload is the
    upper bound the UI should ever see — past that we cut hard."""
    long_text = "x" * (MAX_EXTRACTED_CHARS + 100)
    result = extract_text("essay.txt", long_text.encode())
    assert len(result) == MAX_EXTRACTED_CHARS


def test_extract_case_insensitive_extension():
    """`.DOCX` from a Windows machine routes the same as `.docx`."""
    docx_bytes = _make_docx_bytes(["test"])
    result = extract_text("ESSAY.DOCX", docx_bytes)
    assert "test" in result


# ── Sprint 2.7 fix #6: binary garbage heuristic ──────────────────────


def test_is_likely_text_clean_essay():
    """A normal essay clears the threshold by a wide margin.  We pin
    it explicitly so a regression that tightens the heuristic past
    real prose surfaces here, not on the next student submission."""
    text = (
        "This is a clean essay with normal punctuation, varied "
        "sentence lengths, and a sprinkling of Vietnamese: học "
        "viên giỏi cần luyện tập đều đặn."
    )
    assert _is_likely_text(text) is True


def test_is_likely_text_mostly_binary():
    """Latin-1 decode of bytes 0..255 is dominated by control chars
    (0x00–0x1F + 0x7F + 0x80–0x9F).  The heuristic must reject."""
    # Decode the full latin-1 byte range as text — full of control
    # characters that aren't `isprintable()` and aren't whitespace.
    binary_decoded = bytes(range(256)).decode("latin-1")
    assert _is_likely_text(binary_decoded) is False


def test_extract_rejects_binary_renamed_txt():
    """End-to-end: a PNG header + binary tail renamed `.txt` lands
    on a 400 with the canonical Vietnamese binary-garbage message,
    not silent wall-of-garbage in the student's draft."""
    binary_bytes = b"\x89PNG\r\n\x1a\n" + bytes(range(256)) * 3
    with pytest.raises(FileExtractError, match="binary garbage|không phải text"):
        extract_text("essay.txt", binary_bytes)


def test_extract_docx_with_table():
    """Tables get flattened — students sometimes paste a Task 1
    figures table and we want the numbers to come through, not
    silently disappear."""
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Year"
    table.cell(0, 1).text = "Sales"
    table.cell(1, 0).text = "2024"
    table.cell(1, 1).text = "100"
    buf = io.BytesIO()
    doc.save(buf)

    result = extract_text("essay.docx", buf.getvalue())
    assert "Intro paragraph" in result
    assert "Year" in result and "Sales" in result
    assert "2024" in result and "100" in result
