"""Tests for POST /admin/writing/extract-text — Sprint 19.3 independent
grading file upload.

Auth patched; the underlying file_extract_service runs for real on small
in-memory .docx files (it's pure, fast, no DB). Mirrors the auth-gate +
happy-path style of the other admin-writing tests.
"""

from __future__ import annotations

import io
from unittest.mock import AsyncMock, patch

import docx
from fastapi.testclient import TestClient


def _client() -> TestClient:
    from main import app
    return TestClient(app)


_ADMIN_AUTH = {"Authorization": "Bearer fake.admin.jwt"}
_ADMIN_USER = {"id": "00000000-0000-0000-0000-00000000aaaa", "email": "admin@x"}


def _docx_bytes(paragraphs=None, with_table=False) -> bytes:
    d = docx.Document()
    for p in (paragraphs or ["This is a sample IELTS essay paragraph for testing."]):
        d.add_paragraph(p)
    if with_table:
        t = d.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Year"
        t.rows[0].cells[1].text = "Value"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _upload(content: bytes, filename: str, headers=None):
    files = {"file": (filename, content,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    return _client().post("/admin/writing/extract-text", files=files, headers=headers or {})


# ── Auth ──────────────────────────────────────────────────────────────


def test_extract_requires_auth():
    assert _upload(_docx_bytes(), "essay.docx").status_code == 401


# ── Happy path ────────────────────────────────────────────────────────


def test_extract_docx_returns_text_and_metadata():
    body = _docx_bytes(["Some people believe that technology improves education.",
                        "I largely agree with this view for two reasons."])
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(body, "essay.docx", _ADMIN_AUTH)
    assert r.status_code == 200
    data = r.json()
    assert "technology improves education" in data["extracted_text"]
    assert data["word_count"] > 5
    assert data["file_metadata"]["filename"] == "essay.docx"
    assert data["file_metadata"]["size_bytes"] == len(body)
    assert data["warnings"] == []   # clean prose → no warnings


# ── Rejections ────────────────────────────────────────────────────────


def test_extract_rejects_unsupported_extension():
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(b"%PDF-1.4 fake", "essay.pdf", _ADMIN_AUTH)
    assert r.status_code == 400


def test_extract_rejects_oversize_file():
    # >2MB bytes with a .docx name — the service checks size BEFORE parsing.
    big = b"x" * (2 * 1024 * 1024 + 1)
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(big, "huge.docx", _ADMIN_AUTH)
    assert r.status_code == 400
    assert "lớn" in r.json()["detail"].lower() or "MB" in r.json()["detail"]


def test_extract_rejects_corrupt_docx():
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(b"not a real docx", "essay.docx", _ADMIN_AUTH)
    assert r.status_code == 400


# ── Warnings (informational, non-blocking) ────────────────────────────


def test_extract_surfaces_table_warning():
    body = _docx_bytes(["Intro paragraph with enough words to clear the short-text check easily here."],
                       with_table=True)
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(body, "task1.docx", _ADMIN_AUTH)
    assert r.status_code == 200
    # The table flattens to " | "-joined rows → the endpoint flags it.
    assert any("bảng" in w.lower() for w in r.json()["warnings"])


def test_extract_warns_on_very_short_text():
    body = _docx_bytes(["Hi."])
    with patch("routers.admin_writing.require_admin", new=AsyncMock(return_value=_ADMIN_USER)):
        r = _upload(body, "tiny.docx", _ADMIN_AUTH)
    assert r.status_code == 200
    assert any("ngắn" in w.lower() for w in r.json()["warnings"])
