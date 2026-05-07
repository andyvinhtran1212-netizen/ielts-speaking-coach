"""Tests for services.writing_word_exporter (Sprint W3 Phase 2)."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone

import pytest
from docx import Document

from models.writing_feedback import WritingFeedback
from services.writing_word_exporter import (
    build_filename,
    render_essay_to_docx,
)


# ── Fixtures ─────────────────────────────────────────────────────────

def _l1_payload() -> dict:
    return {
        "overallBandScore": 6.0,
        "overallBandScoreSummary": "Bài đạt mức Band 6.0.",
        "keyTakeaways": {
            "strengths": ["Diễn đạt rõ ý"],
            "areasForImprovement": ["Cần đa dạng cấu trúc câu"],
        },
        "criteriaFeedback": {
            "mainCriterion":     {"title": "Task Response",       "explanation": "x", "feedback": "y", "bandScore": 6},
            "coherenceCohesion": {"title": "Coherence & Cohesion", "explanation": "x", "feedback": "y", "bandScore": 6},
            "lexicalResource":   {"title": "Lexical Resource",     "explanation": "x", "feedback": "y", "bandScore": 6},
            "grammaticalRange":  {"title": "Grammatical Range",    "explanation": "x", "feedback": "y", "bandScore": 6},
        },
        "mistakeAnalysis": [
            {"original": "I has been study", "mistakeType": "Grammar", "explanation": "Sai trợ động từ", "suggestion": "I have been studying", "criterion": "Grammatical Range"},
        ],
        "aiContentAnalysis": {"likelihood": 5, "explanation": "Bài tự nhiên."},
        "improvedEssay": "Improved version.",
    }


def _l5_payload() -> dict:
    payload = _l1_payload()
    payload.update({
        "ideaDevelopmentAnalysis": [
            {"paragraph": 2, "originalIdea": "Học sinh nên học gì họ thích.", "issue": "Chưa rõ", "explanation": "Cần ví dụ.", "suggestion": {"instruction": "Thêm ví dụ", "example": "For instance…"}},
        ],
        "coherenceAnalysis": [
            {"location": "Para 2", "issue": "Topic shift", "explanation": "Đột ngột.", "suggestion": {"instruction": "Câu chuyển ý", "example": "However…"}},
        ],
        "counterargumentAnalysis": {"isPresent": False, "feedback": "Chưa có.", "suggestion": "Thêm đoạn đối lập.", "context": {"insertionPoint": "after p2", "reasoning": "Cân bằng."}},
        "lexicalAnalysis": {"wordsToUpgrade": [{"original": "good", "context": "very good", "suggestions": ["proficient"], "category": "Adjective"}]},
        "sentenceStructureAnalysis": {"sentenceUpgrades": [{"original": "I think.", "rewritten": "It is my contention.", "explanation": "Phức hơn."}]},
    })
    return payload


def _build(payload: dict, *, essay="My essay.", code="S001", task="task2") -> tuple[bytes, str]:
    fb = WritingFeedback(**payload)
    return render_essay_to_docx(
        feedback=fb,
        essay_text=essay,
        prompt_text="Sample prompt.",
        task_type=task,
        student_name="Trần Trọng Vinh",
        student_code=code,
    )


def _open(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


# ── Filename ─────────────────────────────────────────────────────────

@pytest.mark.parametrize("task,expected_t", [
    ("task1_academic", "T1"),
    ("task1_general",  "T1"),
    ("task2",          "T2"),
])
def test_filename_uses_correct_task_suffix(task, expected_t):
    name = build_filename(student_code="S001", task_type=task)
    today = datetime.now(timezone.utc).strftime("%Y%m%d")
    assert name == f"S001_{today}_{expected_t}.docx"


def test_filename_sanitises_unsafe_student_code():
    """Slashes, spaces, etc. dropped — nothing breaks Content-Disposition."""
    name = build_filename(student_code="S/001 hack..", task_type="task2")
    assert "/" not in name
    assert " " not in name
    assert ".." not in name
    assert name.endswith(".docx")


def test_filename_falls_back_when_code_empty():
    name = build_filename(student_code="", task_type="task2")
    assert name.startswith("student_")


# ── Output structure ─────────────────────────────────────────────────

def test_l1_doc_opens_and_has_expected_headings():
    docx_bytes, filename = _build(_l1_payload())
    assert filename.endswith(".docx")
    assert len(docx_bytes) > 1000  # non-trivial size

    doc = _open(docx_bytes)
    headings = [
        p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")
    ]
    # Required L1 sections
    assert "Task 2 Analysis" in headings  # H1
    assert "Overall Band Score" in headings
    assert "Key Takeaways" in headings
    assert "Criteria Breakdown" in headings
    assert "Original Essay (Mistakes Highlighted)" in headings
    assert "Detailed Issue Analysis" in headings
    assert "Improved Essay (Band 8.0+)" in headings
    # No Advanced Analysis at L1
    assert "Advanced Analysis" not in headings


def test_l5_doc_includes_all_advanced_subsections():
    """Sprint 2.5.4 renamed:
        Vocabulary Upgrade            → Vocabulary & Collocation
        Idea Development / Data ...   → Idea Development & Coherence
       Coherence & Flow stays as the secondary coherence sub-heading."""
    docx_bytes, _ = _build(_l5_payload())
    doc = _open(docx_bytes)
    headings = [
        p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")
    ]
    assert "Advanced Analysis" in headings
    assert "Vocabulary & Collocation" in headings
    assert "Sentence Structure Analysis" in headings
    assert "Idea Development & Coherence" in headings
    assert "Coherence & Flow" in headings
    assert "Counterargument" in headings


def test_l1_doc_skips_advanced_when_no_data():
    """Defence: even if a sub-section is set to an empty list, the heading
    must still be skipped."""
    payload = _l1_payload()
    payload["lexicalAnalysis"] = {"wordsToUpgrade": []}
    payload["sentenceStructureAnalysis"] = {"sentenceUpgrades": []}
    payload["coherenceAnalysis"] = []
    payload["ideaDevelopmentAnalysis"] = []
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    assert "Advanced Analysis" not in headings


# ── Content sanity ───────────────────────────────────────────────────

def test_doc_contains_overall_band_text():
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "6.0 / 9.0" in full
    assert "Bài đạt mức Band 6.0." in full
    # Vietnamese diacritics survive
    assert "Trần Trọng Vinh" in full


def test_doc_contains_mistake_table_rows():
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    # Find the mistakes table (5 columns: # / Original / Type / Explanation / Suggestion)
    mistake_tables = [t for t in doc.tables if len(t.columns) == 5]
    assert mistake_tables, "5-col mistakes table missing"
    table = mistake_tables[0]
    # 1 header + 1 mistake row
    assert len(table.rows) == 2
    assert table.rows[1].cells[1].text == "I has been study"
    assert table.rows[1].cells[4].text == "I have been studying"


def test_doc_highlights_mistakes_in_essay():
    """Sprint 2.5.4: highlight switched from yellow w:highlight to red-bg
    w:shd shading (#FEE2E2) so the .docx visually matches the HTML/clipboard
    render. Each mistake's `original` substring carries the shaded background."""
    docx_bytes, _ = _build(
        _l1_payload(),
        essay="I has been study for many years now.",
    )
    doc = _open(docx_bytes)
    shaded_runs = [
        run.text for p in doc.paragraphs for run in p.runs
        if _has_red_bg_shading(run)
    ]
    assert "I has been study" in shaded_runs


def test_doc_handles_empty_mistakes_list_without_table():
    payload = _l1_payload()
    payload["mistakeAnalysis"] = []
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    assert "Detailed Issue Analysis" not in headings


# ── Helpers ──────────────────────────────────────────────────────────

def _is_yellow_highlight(run) -> bool:
    """Inspect run XML for w:highlight w:val='yellow' (legacy mark)."""
    rPr = run._r.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rPr is None:
        return False
    hl = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight")
    if hl is None:
        return False
    return hl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "yellow"


def _has_red_bg_shading(run) -> bool:
    """Sprint 2.5.4: essay highlight now uses w:shd background fill,
    not w:highlight, so a paste from .docx into Word/Google Docs preserves
    a colour rather than the named yellow swatch."""
    rPr = run._r.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rPr is None:
        return False
    shd = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
    if shd is None:
        return False
    fill = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
    return (fill or "").upper() == "FEE2E2"


# ── W3.2 Phase 3: threshold colors + native bullets ──────────────────

from services.writing_word_exporter import _band_color, _COLOR_TEAL, _COLOR_AMBER, _COLOR_RED


@pytest.mark.parametrize("score,expected", [
    (8.0, _COLOR_TEAL),   # high → teal
    (7.0, _COLOR_TEAL),   # boundary → teal
    (6.5, _COLOR_AMBER),  # mid → amber
    (5.5, _COLOR_AMBER),  # boundary → amber
    (5.0, _COLOR_RED),    # low → red
    (3.5, _COLOR_RED),    # well below
])
def test_band_color_threshold(score, expected):
    assert _band_color(score) == expected


def _band_run_color(doc):
    """The first run after the 'Overall Band Score' heading carries the
    coloured band text."""
    found_heading = False
    for p in doc.paragraphs:
        if found_heading and p.runs:
            return p.runs[0].font.color.rgb
        if p.style and p.style.name.startswith("Heading") and p.text == "Overall Band Score":
            found_heading = True
    return None


def test_doc_overall_band_always_blue():
    """Sprint 2.5.4: overall band display is always blue (#2563EB) per
    spec — visual emphasis comes from size (72pt) plus the blue token,
    not from threshold-coloured run colour. The threshold colour now
    lives on per-criterion band cells inside the 2×2 grid.

    Pinning the absolute colour rather than the threshold guards
    against a regression where someone wires _band_color back in for
    the overall display.
    """
    from docx.shared import RGBColor as _RGB
    BLUE = _RGB(0x25, 0x63, 0xEB)
    for score in (4.5, 6.0, 7.5):
        payload = _l1_payload()
        payload["overallBandScore"] = score
        docx_bytes, _ = _build(payload)
        doc = _open(docx_bytes)
        # The 72pt band display is the FIRST centred paragraph after the title.
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for p in doc.paragraphs:
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.runs and p.runs[0].font.size and p.runs[0].font.size.pt >= 60:
                assert p.runs[0].font.color.rgb == BLUE, (
                    f"score={score}: expected blue, got {p.runs[0].font.color.rgb}"
                )
                break
        else:
            raise AssertionError(f"score={score}: 72pt band paragraph not found")


def test_doc_takeaways_use_word_bullet_style():
    """Sprint 2.5.4 takeaways layout is a 1×2 table (was 2-row header+data).
    Each cell carries an inline bold heading + bullet-styled items."""
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    # Key Takeaways is the first table — now 1 row × 2 columns.
    takeaways = doc.tables[0]
    assert len(takeaways.rows) == 1, (
        f"Expected 1×2 takeaways layout, got {len(takeaways.rows)} rows"
    )
    bullet_styles = []
    for cell in (takeaways.rows[0].cells[0], takeaways.rows[0].cells[1]):
        for p in cell.paragraphs:
            # The first paragraph in each cell is the header (bold,
            # not bullet-styled). Anything past it carrying text and
            # styled "List Bullet" is a takeaway item.
            if p.text and p.style.name == "List Bullet":
                bullet_styles.append(p.style.name)
    assert bullet_styles, "no bullet-styled takeaway items rendered"
    assert all(s == "List Bullet" for s in bullet_styles)


# ── Phase 1.5c — sentenceStructureAnalysis dual-shape rendering ──────


def test_doc_renders_phase_1_5c_sentence_structure_shape():
    """Phase 1.5c overloads `sentenceStructureAnalysis` with the
    structured shape (summary + complexity + observation +
    common_issues + focus_theme). Word exporter must detect the
    `summary` key and render the structured block — including the
    focus theme heading + practice line — instead of trying to walk
    a non-existent `sentenceUpgrades` list.

    Pinning the focus_theme.title in the doc text guards against a
    regression where someone adds a Phase-1.5c branch but forgets to
    plumb the focus theme through (the theme is the most actionable
    output for the student)."""
    payload = _l1_payload()
    payload["sentenceStructureAnalysis"] = {
        "summary": "Em hay viết câu chạy không có dấu chấm.",
        "common_issues": [
            {"pattern": "Run-on sentence", "count": 4,
             "examples": ["I went home it was late."]},
        ],
        "complexity_indicator": "needs_more_simple",
        "current_essay_observation": "Bài này có 2 câu chạy.",
        "focus_theme": {
            "title": "Tách câu bằng dấu chấm + liên từ",
            "why":   "Em đang chạy câu — cần fix trước khi push complex structures.",
            "this_week_practice": "Viết 5 câu, mỗi câu kết bằng dấu chấm rõ ràng.",
        },
    }
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)

    headings = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    assert "Sentence Structure Analysis" in headings
    assert "Focus this week" in headings  # H4 sub-heading

    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Summary + observation surface in body text.
    assert "Em hay viết câu chạy" in full_text
    assert "Bài này có 2 câu chạy" in full_text
    # Recurring pattern + count rendered.
    assert "Run-on sentence" in full_text
    assert "(4x)" in full_text
    # Focus theme title + practice line both surface.
    assert "Tách câu bằng dấu chấm" in full_text
    assert "Viết 5 câu" in full_text


# ── Sprint 2.5.3 — recurringPatterns rendering + footer ──────────────


def test_doc_renders_recurring_patterns_section():
    """Sprint 2.5.3 plumbs recurringPatterns through the Word exporter
    (was previously dropped on the floor).  Pin the heading + improvements
    + stillRecurring branches so a future regression that strips one of
    them surfaces immediately."""
    payload = _l1_payload()
    payload["recurringPatterns"] = {
        "summary": "Em đã sửa lỗi article nhưng vẫn còn lỗi tense.",
        "improvements": ["Article a/an/the", "Subject-verb agreement"],
        "stillRecurring": ["Past simple vs present perfect", "Run-on sentences"],
    }
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)

    headings = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    assert "Recurring Patterns" in headings
    assert "Đã cải thiện" in headings
    assert "Vẫn lặp lại" in headings

    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Em đã sửa lỗi article" in full
    assert "Article a/an/the" in full
    assert "Past simple vs present perfect" in full

    # Bullets, not raw paragraphs.
    bullet_paras = [p for p in doc.paragraphs
                    if p.style and p.style.name == "List Bullet"
                    and ("Article" in p.text or "Run-on" in p.text)]
    assert len(bullet_paras) >= 2


def test_doc_renders_recurring_patterns_with_string_still_recurring():
    """stillRecurring has been observed as a bare string in early
    payloads — coerce-to-list path must work without crashing."""
    payload = _l1_payload()
    payload["recurringPatterns"] = {
        "stillRecurring": "Tense agreement only",
    }
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Tense agreement only" in full


def test_doc_skips_recurring_patterns_when_null():
    """No recurringPatterns → no heading."""
    docx_bytes, _ = _build(_l1_payload())  # base payload has no recurringPatterns
    doc = _open(docx_bytes)
    headings = [p.text for p in doc.paragraphs
                if p.style and p.style.name.startswith("Heading")]
    assert "Recurring Patterns" not in headings


def test_doc_includes_footer_credit_line():
    """Sprint 2.5.3 footer pins the 'Generated by Aver Learning' credit
    so admin-shared docs can be identified at a glance."""
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Aver Learning" in full


# ── Sprint 2.5.4 — spec structure pins ───────────────────────────────


def test_takeaways_table_is_1x2_layout():
    """Sprint 2.5.4: takeaways became 1×2 (single row, two cells: green |
    yellow). A future refactor that splits header into a separate row
    would break the cell-shading pinning."""
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    takeaways = doc.tables[0]
    assert len(takeaways.rows) == 1
    assert len(takeaways.rows[0].cells) == 2
    assert "Strengths" in takeaways.rows[0].cells[0].text
    assert "Improvement" in takeaways.rows[0].cells[1].text or \
           "Areas for" in takeaways.rows[0].cells[1].text


def test_criteria_grid_is_2x2_with_titles():
    """Sprint 2.5.4 criteria grid: 2 rows × 2 cols, IELTS criterion names
    fill the 4 cells (mapped from the named bundle to the spec's flat shape)."""
    docx_bytes, _ = _build(_l1_payload())
    doc = _open(docx_bytes)
    grid = None
    for t in doc.tables:
        if (len(t.rows) == 2 and len(t.rows[0].cells) == 2
                and "Task Response" in t.rows[0].cells[0].text):
            grid = t
            break
    assert grid is not None, "2×2 criteria grid not found"
    # All four criterion titles surface across the four cells.
    cell_text = " ".join(
        grid.rows[r].cells[c].text
        for r in range(2) for c in range(2)
    )
    assert "Task Response" in cell_text
    assert "Coherence" in cell_text
    assert "Lexical Resource" in cell_text
    assert "Grammatical Range" in cell_text


def test_lexical_upgrade_3col_table():
    """Sprint 2.5.4 vocabulary subsection: 3-column table — Original | → | Upgrade."""
    docx_bytes, _ = _build(_l5_payload())
    doc = _open(docx_bytes)
    lex_table = None
    for t in doc.tables:
        if (len(t.rows) >= 2 and len(t.rows[0].cells) == 3
                and "Original" in t.rows[0].cells[0].text
                and "Upgrade" in t.rows[0].cells[2].text):
            lex_table = t
            break
    assert lex_table is not None, "3-col lexical upgrade table not found"
    # The data row carries the original word + first upgrade.
    data_row_text = " ".join(c.text for c in lex_table.rows[1].cells)
    assert "good" in data_row_text
    assert "proficient" in data_row_text
    assert "→" in data_row_text


def test_essay_highlight_intervals_helper():
    """find_highlight_intervals merges overlapping matches into one
    interval — both the inner substring and the outer phrase coalesce."""
    from services.writing_render import find_highlight_intervals
    essay = "I are happy. I love English language."
    mistakes = [{"original": "I are"}, {"original": "English language"}]
    intervals = find_highlight_intervals(essay, mistakes)
    assert len(intervals) == 2
    assert intervals[0] == (0, 5)
    assert intervals[1][0] == essay.index("English language")


def test_essay_highlight_intervals_skips_short_originals():
    """`original` strings shorter than 3 chars are skipped — single-letter
    mistakes (eg "a" → "an") would otherwise highlight every 'a' in the essay."""
    from services.writing_render import find_highlight_intervals
    essay = "An example sentence."
    mistakes = [{"original": "a"}, {"original": "An"}]
    intervals = find_highlight_intervals(essay, mistakes)
    assert intervals == []


def test_doc_renders_legacy_sentence_upgrades_shape():
    """The legacy `{sentenceUpgrades: [{original, rewritten,
    explanation}]}` shape must keep rendering unchanged — Phase 1.5c
    only fires for ≥5-essay students, so L4/L5 essays from new
    students still emit this shape via the system prompt.

    Phase 1.5c rework relaxed `sentenceStructureAnalysis` from a
    Pydantic class to `Optional[dict]`, so this test also pins the
    dict-access path (.get('sentenceUpgrades') vs .sentenceUpgrades)
    against a future regression where someone re-introduces attribute
    access without updating the field type."""
    payload = _l5_payload()
    # _l5_payload already provides legacy sentenceUpgrades — assert it
    # surfaces as before.
    docx_bytes, _ = _build(payload)
    doc = _open(docx_bytes)

    headings = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    assert "Sentence Structure Analysis" in headings
    # Phase-1.5c-only sub-heading must NOT appear when shape is legacy.
    assert "Focus this week" not in headings

    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Original (italicised) + rewritten + explanation all surface.
    assert "I think." in full_text
    assert "It is my contention." in full_text
    assert "Phức hơn." in full_text
