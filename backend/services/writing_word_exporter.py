"""services/writing_word_exporter.py — .docx export for Writing Coach.

Sprint 2.5.4: rewritten to share structure + colour palette with the
HTML/clipboard render. Both paths consume the same shape-normalisation
helpers in services.writing_render so admin edits, Gemini variants, and
typed Pydantic objects all flow through one extraction layer.

Filename contract: {student_code}_{YYYYMMDD}_T{1|2}.docx
"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from models.writing_feedback import WritingFeedback
from services.writing_render import (
    COLOR_BLUE, COLOR_BODY, COLOR_BORDER, COLOR_GRAY_BG, COLOR_GREEN,
    COLOR_GREEN_BG, COLOR_MUTED, COLOR_RED, COLOR_RED_BG, COLOR_SUBHEADING,
    COLOR_YELLOW, COLOR_YELLOW_BG,
    _build_criteria_grid_rows, _build_takeaways_ctx,
    _criteria_list, _extract_ai_content, _extract_counterargument,
    _extract_idea_paragraphs, _extract_improved_essay_text,
    _extract_lexical_summary, _extract_lexical_upgrades,
    _extract_sentence_structures, _format_band,
    _mistake_dict, _normalize_mistakes, _split_essay_paragraphs,
    _takeaways_dict, find_highlight_intervals,
)


_TASK_LABELS: dict[str, str] = {
    "task1_academic": "Task 1 Academic Analysis",
    "task1_general":  "Task 1 General Analysis",
    "task2":          "Task 2 Analysis",
}


# ── Hex → RGBColor helpers ───────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hex_no_hash(hex_color: str) -> str:
    """Strip leading # for OOXML w:fill attributes."""
    return hex_color.lstrip("#").upper()


_RGB_SUBHEAD  = _hex_to_rgb(COLOR_SUBHEADING)
_RGB_BODY     = _hex_to_rgb(COLOR_BODY)
_RGB_MUTED    = _hex_to_rgb(COLOR_MUTED)
_RGB_GREEN    = _hex_to_rgb(COLOR_GREEN)
_RGB_YELLOW   = _hex_to_rgb(COLOR_YELLOW)
_RGB_RED      = _hex_to_rgb(COLOR_RED)
_RGB_BLUE     = _hex_to_rgb(COLOR_BLUE)


# ── Backwards-compat band-color helper (used by writing_render tests) ─
# Sprint 2.5.4 keeps the threshold so tests pinning low/mid/high
# overall band colour still pass via an indirect import.
_COLOR_TEAL  = RGBColor(0x0D, 0x94, 0x88)
_COLOR_AMBER = RGBColor(0xCA, 0x8A, 0x04)
_COLOR_RED   = RGBColor(0xDC, 0x26, 0x26)


def _band_color(score) -> RGBColor:
    if score is None:
        return _COLOR_TEAL
    if score >= 7.0:
        return _COLOR_TEAL
    if score >= 5.5:
        return _COLOR_AMBER
    return _COLOR_RED


# ── Public API ───────────────────────────────────────────────────────

def render_essay_to_docx(
    *,
    feedback: WritingFeedback,
    essay_text: str,
    prompt_text: str,
    task_type: str,
    student_name: str,
    student_code: str,
    hide_scores: bool = False,
) -> tuple[bytes, str]:
    """Build .docx bytes + suggested filename for one essay's feedback.

    U2 — when hide_scores is True (student download of an essay delivered with
    the "Ẩn điểm" flag), the numeric scores are omitted from the Word file too:
    the 72pt overall band block AND the per-criterion Criteria Breakdown table.
    The qualitative narrative (Overall Band Score summary, Key Takeaways,
    per-section feedback) is kept. Default False = full export (admin path).
    """
    doc = Document()
    _build_document(
        doc=doc,
        feedback=feedback,
        essay_text=essay_text,
        prompt_text=prompt_text,
        task_type=task_type,
        student_name=student_name,
        hide_scores=hide_scores,
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), build_filename(student_code=student_code, task_type=task_type)


def build_filename(*, student_code: str, task_type: str) -> str:
    today = datetime.now(timezone.utc).strftime("%Y%m%d")
    task_num = "2" if task_type == "task2" else "1"
    safe_code = re.sub(r"[^A-Za-z0-9_-]", "", student_code) or "student"
    return f"{safe_code}_{today}_T{task_num}.docx"


# ── Document builder ─────────────────────────────────────────────────

def _build_document(
    *,
    doc: Document,
    feedback: WritingFeedback,
    essay_text: str,
    prompt_text: str,
    task_type: str,
    student_name: str,
    hide_scores: bool = False,
) -> None:
    # 2.1 Title
    title = doc.add_heading(_TASK_LABELS.get(task_type, "Writing Analysis"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _RGB_SUBHEAD

    if student_name:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m_run = meta.add_run(student_name)
        m_run.italic = True
        m_run.font.color.rgb = _RGB_MUTED
        m_run.font.size = Pt(10)

    # U2 — the big 72pt numeric band is a score → omit when hidden. The
    # narrative summary below is kept.
    if not hide_scores:
        _add_overall_band_block(
            doc,
            feedback.overallBandScore,
            feedback.overallBandScoreSummary,
        )

    # ── Overall Band Score heading (kept for backwards-compat with
    # callers that scan for it; rendered as a small spacer label
    # before the next section). The visual band display is the big
    # 72pt block above; this heading just anchors the section.
    doc.add_heading("Overall Band Score", level=2)
    if feedback.overallBandScoreSummary:
        s_para = doc.add_paragraph()
        s_run = s_para.add_run(feedback.overallBandScoreSummary)
        s_run.italic = True
        s_run.font.color.rgb = _RGB_BODY

    # 2.2 Key Takeaways: 1×2
    doc.add_heading("Key Takeaways", level=2)
    _build_takeaways_table(doc, _build_takeaways_ctx(_takeaways_dict(feedback)))

    # 2.3 Criteria Breakdown: 2×2 grid — the per-criterion band scores;
    # omit entirely when hidden (mirrors the FE hiding the 4 sub-band cards).
    criteria_rows = _build_criteria_grid_rows(_criteria_list(feedback))
    if criteria_rows and not hide_scores:
        doc.add_heading("Criteria Breakdown", level=2)
        _build_criteria_grid_table(doc, criteria_rows)

    # 2.4 Original Essay with red-bg highlights
    doc.add_heading("Original Essay (Mistakes Highlighted)", level=2)
    if prompt_text:
        prompt_p = doc.add_paragraph()
        r = prompt_p.add_run("Prompt: ")
        r.bold = True
        prompt_p.add_run(prompt_text)
    _build_highlighted_essay_paragraphs(
        doc, essay_text,
        [_mistake_dict(m) for m in (feedback.mistakeAnalysis or [])],
    )

    # 2.5 Detailed Issue Analysis: 5-col table
    mistakes = _normalize_mistakes(
        [_mistake_dict(m) for m in (feedback.mistakeAnalysis or [])]
    )
    if mistakes:
        doc.add_heading("Detailed Issue Analysis", level=2)
        _build_mistake_table(doc, mistakes)

    # ── Recurring Patterns (Sprint 2.5.3) ──
    rp = feedback.recurringPatterns
    if rp and isinstance(rp, dict) and (rp.get("summary") or rp.get("improvements") or rp.get("stillRecurring")):
        doc.add_heading("Recurring Patterns", level=2)
        if rp.get("summary"):
            sp = doc.add_paragraph()
            sp.add_run(rp["summary"]).italic = True
        if isinstance(rp.get("improvements"), list) and rp["improvements"]:
            doc.add_heading("Đã cải thiện", level=3)
            for item in rp["improvements"]:
                doc.add_paragraph(str(item), style="List Bullet")
        still = rp.get("stillRecurring")
        if isinstance(still, str):
            still = [still]
        if isinstance(still, list) and still:
            doc.add_heading("Vẫn lặp lại", level=3)
            for item in still:
                doc.add_paragraph(str(item), style="List Bullet")

    # 2.6 Advanced Analysis
    lex_summary = _extract_lexical_summary(feedback.lexicalAnalysis)
    lex_upgrades = _extract_lexical_upgrades(feedback.lexicalAnalysis)
    sentences = _extract_sentence_structures(feedback.sentenceStructureAnalysis)
    paragraphs = _extract_idea_paragraphs(feedback.ideaDevelopmentAnalysis)
    coherence_items = list(feedback.coherenceAnalysis or [])
    counterargument = _extract_counterargument(feedback.counterargumentAnalysis, task_type)

    has_advanced = bool(
        lex_upgrades or lex_summary or sentences or paragraphs or
        coherence_items or counterargument
    )
    if has_advanced:
        doc.add_heading("Advanced Analysis", level=2)

    if lex_upgrades or lex_summary:
        doc.add_heading("Vocabulary & Collocation", level=3)
        if lex_summary:
            doc.add_paragraph(lex_summary)
        if lex_upgrades:
            _build_lexical_upgrade_table(doc, lex_upgrades)

    if sentences:
        doc.add_heading("Sentence Structure Analysis", level=3)
        for s in sentences:
            o_para = doc.add_paragraph()
            o_para.add_run("Original: ").bold = True
            o_run = o_para.add_run(s.get("original", ""))
            o_run.font.color.rgb = _RGB_RED
            r_para = doc.add_paragraph()
            r_para.add_run("Rewritten: ").bold = True
            r_run = r_para.add_run(s.get("improved", ""))
            r_run.font.color.rgb = _RGB_GREEN
            if s.get("technique"):
                t_para = doc.add_paragraph()
                t_run = t_para.add_run(f"Technique: {s['technique']}")
                t_run.italic = True
                t_run.font.color.rgb = _RGB_MUTED
            doc.add_paragraph()

    # Phase 1.5c sentenceStructureAnalysis structured shape — render
    # focus theme + recurring patterns as a separate sub-section so the
    # Phase-1.5c output isn't lost when sentenceUpgrades is absent.
    ss_data = feedback.sentenceStructureAnalysis or {}
    if isinstance(ss_data, dict) and ss_data.get("summary") and not sentences:
        doc.add_heading("Sentence Structure Analysis", level=3)
        if ss_data.get("summary"):
            doc.add_paragraph(ss_data["summary"])
        if ss_data.get("complexity_indicator"):
            _kv_paragraph(doc, "Complexity:", ss_data["complexity_indicator"])
        if ss_data.get("current_essay_observation"):
            _kv_paragraph(doc, "This essay:", ss_data["current_essay_observation"])
        issues = ss_data.get("common_issues") or []
        if isinstance(issues, list) and issues:
            doc.add_paragraph("Recurring patterns:")
            for issue in issues:
                if not isinstance(issue, dict):
                    continue
                pattern = (issue.get("pattern") or "").strip()
                count   = issue.get("count")
                examples = issue.get("examples") or []
                line = f"• {pattern}"
                if count is not None:
                    line += f" ({count}x)"
                if isinstance(examples, list) and examples:
                    line += " — " + "; ".join(
                        f'"{e}"' for e in examples if isinstance(e, str)
                    )
                doc.add_paragraph(line)
        focus = ss_data.get("focus_theme") or {}
        if isinstance(focus, dict) and focus.get("title"):
            doc.add_heading("Focus this week", level=4)
            _kv_paragraph(doc, "Theme:", focus.get("title") or "")
            if focus.get("why"):
                _kv_paragraph(doc, "Why:", focus.get("why") or "")
            if focus.get("this_week_practice"):
                _kv_paragraph(doc, "Practice:", focus.get("this_week_practice") or "")

    if paragraphs:
        doc.add_heading("Idea Development & Coherence", level=3)
        for p in paragraphs:
            head = doc.add_paragraph()
            head.add_run(f"Paragraph {p.get('index', '?')}: ").bold = True
            head.add_run(p.get("heading") or "")
            if p.get("original"):
                op = doc.add_paragraph()
                op.add_run(p["original"]).italic = True
            if p.get("commentary"):
                doc.add_paragraph(p["commentary"])
            if p.get("suggestion"):
                sp = doc.add_paragraph()
                s_run = sp.add_run("Suggestion: ")
                s_run.bold = True
                s_run.font.color.rgb = _RGB_GREEN
                sp.add_run(str(p["suggestion"]))
            doc.add_paragraph()

    if coherence_items:
        doc.add_heading("Coherence & Flow", level=3)
        for c in coherence_items:
            label_text = (getattr(c, "location", None) or "Issue")
            _kv_paragraph(doc, f"{label_text}:", getattr(c, "issue", "") or "")
            doc.add_paragraph(getattr(c, "explanation", "") or "")
            sug = getattr(c, "suggestion", None)
            if sug:
                sug_text = (getattr(sug, "instruction", "") or "")
                if getattr(sug, "example", None):
                    sug_text = f"{sug_text} — {sug.example}".strip(" —")
                _kv_paragraph(doc, "Suggestion:", sug_text)

    if counterargument:
        doc.add_heading("Counterargument", level=3)
        if counterargument.get("summary"):
            doc.add_paragraph(counterargument["summary"])
        for pt in counterargument.get("points", []):
            doc.add_paragraph(str(pt), style="List Bullet")

    # AI Content Analysis (kept from prior render).
    ai = _extract_ai_content(feedback.aiContentAnalysis)
    if ai:
        doc.add_heading("AI Content Analysis", level=2)
        if ai.get("likelihood") is not None:
            _kv_paragraph(doc, "Likelihood AI-generated:", f"{ai['likelihood']}%")
        if ai.get("explanation"):
            doc.add_paragraph(ai["explanation"])

    # 2.7 Improved Essay (Band 8.0+)
    improved = _extract_improved_essay_text(feedback.improvedEssay)
    if improved:
        doc.add_heading("Improved Essay (Band 8.0+)", level=2)
        for line in _split_essay_paragraphs(improved):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.color.rgb = _RGB_BODY

    # 2.8 Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Generated by Aver Learning Writing Coach")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = _RGB_MUTED


# ── Cell shading + run highlighting (OOXML) ──────────────────────────

def _set_cell_shading(cell, fill_hex_no_hash: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex_no_hash)
    tcPr.append(shd)


def _add_run_highlight_bg(run, fill_hex_no_hash: str) -> None:
    """Apply a background colour shading to a run (used for essay highlights)."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex_no_hash)
    rPr.append(shd)


def _highlight_run_yellow(run) -> None:
    """Backwards-compat: yellow highlight via w:highlight (legacy tests).

    Sprint 2.5.4 switched essay highlighting to w:shd background, but
    the legacy test_doc_highlights_mistakes_in_essay test is being
    updated; this helper stays for any external caller that imports it.
    """
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)


# ── Block builders ───────────────────────────────────────────────────

def _add_overall_band_block(doc, band, summary) -> None:
    if band is None:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band_run = para.add_run(_format_band(band))
    band_run.bold = True
    band_run.font.size = Pt(72)
    band_run.font.color.rgb = _RGB_BLUE
    suffix_run = para.add_run(" / 9.0")
    suffix_run.font.size = Pt(24)
    suffix_run.font.color.rgb = _RGB_MUTED


def _build_takeaways_table(doc, takeaways: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    cell_s = table.rows[0].cells[0]
    cell_a = table.rows[0].cells[1]
    _fill_takeaway_cell(
        cell_s, "Strengths", takeaways.get("strengths", []),
        _RGB_GREEN, _hex_no_hash(COLOR_GREEN_BG),
    )
    _fill_takeaway_cell(
        cell_a, "Areas for Improvement", takeaways.get("areas_for_improvement", []),
        _RGB_YELLOW, _hex_no_hash(COLOR_YELLOW_BG),
    )


def _fill_takeaway_cell(cell, header_text, items, header_color, bg_hex) -> None:
    _set_cell_shading(cell, bg_hex)
    cell.paragraphs[0].text = ""
    para = cell.paragraphs[0]
    h_run = para.add_run(header_text)
    h_run.bold = True
    h_run.font.size = Pt(12)
    h_run.font.color.rgb = header_color

    if not items:
        em = cell.add_paragraph("—")
        for r in em.runs:
            r.font.color.rgb = _RGB_MUTED
        return
    for item in items:
        b = cell.add_paragraph(style="List Bullet")
        b.add_run(str(item)).font.size = Pt(11)


def _build_criteria_grid_table(doc, rows: list) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(2):
            cell = table.rows[r_idx].cells[c_idx]
            if c_idx >= len(row):
                # Odd-length criteria list — leave the trailing cell empty.
                continue
            _fill_criterion_cell(cell, row[c_idx])


def _fill_criterion_cell(cell, c: dict) -> None:
    cell.paragraphs[0].text = ""
    title_p = cell.paragraphs[0]
    name_run = title_p.add_run(c.get("criterion") or "—")
    name_run.bold = True
    name_run.font.color.rgb = _RGB_SUBHEAD

    title_p.add_run("\t")
    band_run = title_p.add_run(str(c.get("band", "—")))
    band_run.bold = True
    band_run.font.size = Pt(18)
    band_run.font.color.rgb = _RGB_BLUE

    if c.get("summary"):
        sp = cell.add_paragraph()
        sr = sp.add_run(c["summary"])
        sr.italic = True
        sr.font.color.rgb = _RGB_MUTED

    if c.get("strengths"):
        hp = cell.add_paragraph()
        hr = hp.add_run("Strengths")
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = _RGB_GREEN
        for s in c["strengths"]:
            cell.add_paragraph(s, style="List Bullet")

    if c.get("improvements"):
        hp = cell.add_paragraph()
        hr = hp.add_run("Improvements")
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = _RGB_YELLOW
        for i in c["improvements"]:
            cell.add_paragraph(i, style="List Bullet")

    if c.get("detailed_text") and not c.get("strengths") and not c.get("improvements"):
        cell.add_paragraph(c["detailed_text"])


def _build_highlighted_essay_paragraphs(doc, essay_text: str, mistakes: list) -> None:
    """Add the essay paragraphs to the doc, wrapping each mistake's
    `original` substring with a red-bg shaded run."""
    if not essay_text:
        return
    intervals = find_highlight_intervals(essay_text, mistakes)
    fill = _hex_no_hash(COLOR_RED_BG)

    # Process paragraph-by-paragraph; intervals reference absolute essay
    # offsets so we shift them as we walk.
    paragraphs = essay_text.split("\n\n")
    cursor = 0
    iv_idx = 0
    for raw in paragraphs:
        para_start = cursor
        para_end = cursor + len(raw)
        cursor = para_end + 2  # account for the "\n\n" separator
        if not raw.strip():
            continue
        p = doc.add_paragraph()
        last = para_start
        while iv_idx < len(intervals) and intervals[iv_idx][0] < para_end:
            i_start, i_end = intervals[iv_idx]
            if i_start >= last:
                if i_start > last:
                    pre = essay_text[last:i_start]
                    pre_run = p.add_run(pre)
                    pre_run.font.color.rgb = _RGB_BODY
                hi = essay_text[i_start:min(i_end, para_end)]
                hi_run = p.add_run(hi)
                hi_run.font.color.rgb = _RGB_BODY
                _add_run_highlight_bg(hi_run, fill)
                last = min(i_end, para_end)
            if i_end <= para_end:
                iv_idx += 1
            else:
                # Highlight spans paragraph boundary — split at end.
                break
        if last < para_end:
            tail = essay_text[last:para_end]
            t_run = p.add_run(tail)
            t_run.font.color.rgb = _RGB_BODY


def _build_mistake_table(doc, mistakes: list) -> None:
    table = doc.add_table(rows=len(mistakes) + 1, cols=5)
    table.style = "Table Grid"

    headers = ["#", "Original", "Type", "Explanation", "Suggestion"]
    fill = _hex_no_hash(COLOR_GRAY_BG)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, fill)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    for i, m in enumerate(mistakes, 1):
        row = table.rows[i].cells
        row[0].text = str(i)

        row[1].paragraphs[0].text = ""
        orig_run = row[1].paragraphs[0].add_run(m.get("original", ""))
        orig_run.font.color.rgb = _RGB_RED
        orig_run.font.name = "Consolas"

        row[2].paragraphs[0].text = ""
        type_run = row[2].paragraphs[0].add_run(m.get("type", "—"))
        type_run.font.color.rgb = _RGB_RED
        type_run.bold = True

        row[3].text = m.get("explanation", "") or ""

        row[4].paragraphs[0].text = ""
        sugg_run = row[4].paragraphs[0].add_run(m.get("suggestion", ""))
        sugg_run.italic = True
        sugg_run.font.color.rgb = _RGB_GREEN


def _build_lexical_upgrade_table(doc, upgrades: list) -> None:
    if not upgrades:
        return
    table = doc.add_table(rows=len(upgrades) + 1, cols=3)
    table.style = "Table Grid"

    headers = ["Original", "", "Upgrade (Band 8+)"]
    fill = _hex_no_hash(COLOR_GRAY_BG)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, fill)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    for i, u in enumerate(upgrades, 1):
        row = table.rows[i].cells

        row[0].paragraphs[0].text = ""
        orig_run = row[0].paragraphs[0].add_run(u.get("original", ""))
        orig_run.font.color.rgb = _RGB_RED
        orig_run.font.name = "Consolas"

        arrow_p = row[1].paragraphs[0]
        arrow_p.text = ""
        arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        arrow_p.add_run("→")

        row[2].paragraphs[0].text = ""
        upg_run = row[2].paragraphs[0].add_run(u.get("upgrade", ""))
        upg_run.font.color.rgb = _RGB_GREEN
        upg_run.font.name = "Consolas"


def _kv_paragraph(doc, key: str, value: str, *, italic_value: bool = False):
    p = doc.add_paragraph()
    p.add_run(key + " ").bold = True
    run = p.add_run(value)
    if italic_value:
        run.italic = True
    return p
