"""services/file_extract_service.py — Phase 2.3c-2 essay-file parser.

Extracts plain text from a student-uploaded essay file so the submit
form can append it to the textarea.  The endpoint that calls into here
(`POST /api/writing/extract-text`) is intentionally stateless — the
textarea remains the single source of truth for draft state, so a
parse failure can't corrupt the saved draft.

Format scope (Phase 2.3c-2):
  •  .docx — `python-docx` (already pinned at 1.1.2 for the Word
     export pipeline; we reuse it).  Paragraphs joined with blank
     lines; tables flattened into ` | `-delimited rows so simple
     two-column "task description / chart figures" tables don't get
     dropped.
  •  .txt  — try a fallback chain of encodings (UTF-8, UTF-8 BOM,
     latin-1, cp1252).  Vietnamese students export from a mix of
     Word/Pages/Notepad and the trailing fallbacks save us from a
     stray cp1252 file blowing up at decode time.

Size cap: 2 MB.  Plenty of headroom for an essay (10k chars × 2 bytes
= 20 KB) — the 2 MB cap is really there to swat away PDFs renamed to
`.docx` and copy-pasted screenshot blobs.

Char cap: 15 000.  The Pydantic `essay_text` validator on submit caps
at 10 k; we leave 5 k of slack so the UI can show a truncation hint
without rejecting outright.  `MAX_EXTRACTED_CHARS` is the contract
the test suite pins.
"""

from __future__ import annotations

import io
import logging
from typing import Optional

import docx  # python-docx — same dep used by services/word_export.py

logger = logging.getLogger(__name__)


MAX_FILE_SIZE_BYTES = 2 * 1024 * 1024
ALLOWED_EXTENSIONS = {".docx", ".txt"}
MAX_EXTRACTED_CHARS = 15_000


class FileExtractError(Exception):
    """Raised on any user-facing extraction failure (bad file, bad
    encoding, oversize, empty).  The router maps every instance to a
    400 with the original Vietnamese message intact."""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Concatenate paragraphs + flattened table rows from a .docx.

    Falls over with `FileExtractError` if python-docx can't open the
    bytes — the SDK raises a mix of `PackageNotFoundError`,
    `BadZipFile`, and the occasional plain `Exception`, so we catch
    broadly here and surface a single user-readable message instead
    of leaking the SDK internals."""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise FileExtractError(f"Không đọc được file .docx: {exc}")

    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables flattened best-effort — we don't try to preserve column
    # alignment, just rescue any text that lives inside a table cell.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts).strip()


# Sprint 2.7 fix #6: binary-detection threshold for the printable
# heuristic.  0.85 lets through clean prose (typical essays score
# 0.99+) and Vietnamese diacritics (all printable Unicode letters)
# while rejecting renamed PNGs / PDFs whose latin-1 decode is full
# of control bytes (0x00–0x1F + 0x7F + 0x80–0x9F).  Tuning lower
# than 0.85 starts admitting binary garbage; higher than 0.90
# starts catching legitimate poetry with heavy whitespace
# formatting.
_TXT_PRINTABLE_THRESHOLD = 0.85


def _is_likely_text(decoded: str, threshold: float = _TXT_PRINTABLE_THRESHOLD) -> bool:
    """Return True when at least `threshold` fraction of characters
    are printable or common whitespace.

    Why we need this: latin-1 decode never raises (it's a 1-to-1 byte
    map), so a binary file renamed `.txt` would otherwise sail
    through the fallback chain and land in the textarea as a
    payload of control characters.  The student then sees garbage
    and the grader gets a corrupted essay row.

    `str.isprintable()` returns False for control characters (Cc),
    formatting marks (Cf), and surrogates (Cs); whitelisting
    `\\n \\r \\t` keeps real text files from tripping the heuristic
    on their line breaks.
    """
    if not decoded:
        return True  # empty is handled by the size/empty checks elsewhere

    printable_count = sum(
        1 for c in decoded
        if c.isprintable() or c in "\n\r\t"
    )
    return (printable_count / len(decoded)) >= threshold


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Try a fallback chain of encodings, then validate the decoded
    output looks like text.

    UTF-8 (incl. BOM) is tried first because that's what Word /
    Pages / VS Code emit by default; latin-1 / cp1252 cover the
    common Vietnamese-Windows exports.  latin-1 will always succeed
    on any byte stream, so it's the implicit safety net for the
    decode step — but a successful decode of a binary file is still
    semantic garbage, which is what `_is_likely_text` catches.

    Sprint 2.7 fix #6: the printable-ratio heuristic AFTER decode.
    A renamed PNG/PDF/exe would previously sail through latin-1 and
    paste a wall of control characters into the student's draft;
    now it raises a 400 with the canonical Vietnamese message.
    """
    encodings = ["utf-8-sig", "utf-8", "latin-1", "cp1252"]
    decoded: Optional[str] = None
    for enc in encodings:
        try:
            decoded = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue

    if decoded is None:
        raise FileExtractError("Không đọc được file .txt — encoding không hỗ trợ.")

    if not _is_likely_text(decoded):
        raise FileExtractError(
            "File .txt chứa nội dung không phải text (binary garbage). "
            "Vui lòng upload file text thuần."
        )

    return decoded.strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Validate then dispatch to the right extractor.

    Validation order matters: we check size BEFORE extension so a
    100 MB renamed PDF gets the friendlier "quá lớn" message and we
    never spend memory parsing it.
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileExtractError(
            f"File quá lớn. Tối đa {MAX_FILE_SIZE_BYTES // 1024 // 1024}MB."
        )
    if len(file_bytes) == 0:
        raise FileExtractError("File rỗng.")

    fname_lower = (filename or "").lower()

    if fname_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif fname_lower.endswith(".txt"):
        text = extract_text_from_txt(file_bytes)
    else:
        raise FileExtractError(
            "Định dạng không hỗ trợ. Chỉ chấp nhận: "
            + ", ".join(sorted(ALLOWED_EXTENSIONS))
        )

    if not text:
        raise FileExtractError("File không có nội dung text.")

    if len(text) > MAX_EXTRACTED_CHARS:
        logger.info(
            "file_extract truncated: orig=%d → max=%d",
            len(text), MAX_EXTRACTED_CHARS,
        )
        text = text[:MAX_EXTRACTED_CHARS]

    return text
