"""services/listening_convert.py — Sprint 13.4 (DEBT-ADMIN-LISTENING-AUTHORING 6/N).

Parse Andy's 2-file Cambridge IELTS DOCX bundle (Question Paper + Script+
AnswerKey) into structured test data for ingest into ``listening_tests``
+ ``listening_content`` + ``listening_exercises``.

Public entry point: ``parse_listening_test(qp_bytes, sa_bytes)``.

All functions are pure (no DB / no network). The router orchestrates I/O.

Sample format (Pilot 01 — ILR-LIS-001):
  * Question Paper carries the student-facing prompts (Q1-Q40) broken by
    section. Each section starts with an instruction line ("Complete the
    form ...") that selects a question-type discriminator.
  * Script+AnswerKey carries:
      - a metadata table at the top (Test ID, Version, Band Target, etc.)
      - four "SECTION N — Transcript" blocks with speaker tags + delivery
        cues + ``(Q##)`` markers embedded
      - a per-section answer key table (Question | Answer | Trap mechanism)

Marker grammar (stripped from user-facing transcript, preserved in
``metadata.raw_transcript``):

* Speaker tags         ``[F-BrE-30s-professional]``  (gender + accent + age + register)
* Delivery cues        ``[pace:slow]`` ``[pause:2s]`` ``[emphasis:word]`` ``[stress:word]``
                       ``[emotion:concerned]`` ``[hesitation:um]`` ``[chuckle:soft]``
* Self-closing cues    ``[hesitate]`` ``[breath]`` ``[sigh]`` ``[chuckle]``
* Question pointers    ``(Q1)`` ``(Q11)`` ``(Q21)`` ``(Q31)``

The router/UI cap file size at 5 MB per DOCX (Sprint 13.4 §Style).
"""

from __future__ import annotations

import io
import re
from typing import Any

try:                                            # python-docx is in requirements;
    from docx import Document                   # the import-time fail-soft keeps
    _DOCX_IMPORT_ERROR: Exception | None = None # unit tests that bypass DOCX
except Exception as exc:                        # IO unaffected.
    Document = None                             # type: ignore[assignment]
    _DOCX_IMPORT_ERROR = exc


# ── Marker regexes ──────────────────────────────────────────────────────────


# [F-BrE-30s-professional], [M-AusE-40s-casual], [N-BrE].
# Permissive whitespace; gender optional; age/register optional.
_SPEAKER_TAG_RE = re.compile(
    r"\[\s*([MFN])\s*-\s*([A-Za-z]+)"
    r"(?:\s*-\s*([0-9]+s|teens|adult))?"
    r"(?:\s*-\s*([A-Za-z][A-Za-z\-]*))?"
    r"\s*\]",
    re.IGNORECASE,
)

# Bracket cues with a colon payload: pace:slow / pause:2s / emphasis:word ...
_DELIVERY_CUE_RE = re.compile(
    r"\[\s*(?:pace|pause|emphasis|stress|emotion|hesitation|chuckle)"
    r"\s*:\s*[^\]]+\]",
    re.IGNORECASE,
)

# Self-closing bracket flags with no colon payload.
_FLAG_CUE_RE = re.compile(
    r"\[\s*(?:hesitate|breath|sigh|chuckle)\s*\]",
    re.IGNORECASE,
)

# (Q1), (Q11), ( Q 33 ) — case + whitespace permissive.
_QUESTION_MARKER_RE = re.compile(r"\(\s*Q\s*(\d{1,2})\s*\)", re.IGNORECASE)

# "1 ............" / "1. ………" / "1\t........." — leading question number.
# Unicode ellipsis variants (… vs ... vs ………) are all valid Cambridge.
_QUESTION_NUMBER_LINE_RE = re.compile(
    r"^\s*(\d{1,2})\s*[\.\)]?\s+(.+?)\s*$",
)

# Section split (Script+AnswerKey): "SECTION 1 — Transcript", "Section 1 - Transcript"
_SECTION_HEADER_RE = re.compile(
    r"^\s*SECTION\s+([1-4])\s*[—–\-]\s*Transcript\s*$",
    re.IGNORECASE,
)

# Section instruction (Question Paper) → question type:
#   "Complete the form" / "Complete the notes" / "Complete the table"
#   "Complete the sentences" → dictation_gap_fill
#   "Answer the questions"  → dictation_short_answer
#   "Choose the correct letter, A, B or C" → mcq_3option
#   "Label the plan/diagram/map" → mcq_letter_label (A-H)
_INSTRUCTION_HINTS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"label the (?:plan|diagram|map)", re.IGNORECASE),
        "mcq_letter_label"),
    (re.compile(r"choose the correct letter", re.IGNORECASE),
        "mcq_3option"),
    (re.compile(r"answer the questions", re.IGNORECASE),
        "dictation_short_answer"),
    (re.compile(r"complete the (?:form|notes?|table|sentences?)", re.IGNORECASE),
        "dictation_gap_fill"),
]


# ── DOCX extraction ─────────────────────────────────────────────────────────


def _extract_docx(file_bytes: bytes) -> tuple[str, list[list[list[str]]]]:
    """Read a DOCX byte stream and return ``(full_text, tables)``.

    ``full_text`` joins paragraphs with ``\n``. ``tables`` is a list of
    tables; each table is a 2D list of cell strings.
    """
    if Document is None:                        # pragma: no cover — install guard
        raise RuntimeError(
            f"python-docx not available: {_DOCX_IMPORT_ERROR}"
        )
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return "\n".join(paragraphs), tables


# ── Marker stripping ────────────────────────────────────────────────────────


def strip_markers(raw: str) -> str:
    """Strip all known markers, returning user-facing transcript text.

    Order matters: speaker tags first (longest pattern), then delivery
    cues, then self-closing flags, then question pointers. Collapse
    repeated whitespace and trim leading/trailing space per line.
    """
    out = _SPEAKER_TAG_RE.sub(" ", raw)
    out = _DELIVERY_CUE_RE.sub(" ", out)
    out = _FLAG_CUE_RE.sub(" ", out)
    out = _QUESTION_MARKER_RE.sub(" ", out)
    # Collapse runs of horizontal whitespace inside lines, preserve \n.
    out = re.sub(r"[ \t]+", " ", out)
    out = re.sub(r" *\n *", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def parse_speakers(raw: str) -> list[dict[str, str | None]]:
    """Return de-duplicated speaker structs in order of first appearance.

    Each struct: ``{tag, gender, accent, age, register}``.
    """
    seen: dict[str, dict[str, str | None]] = {}
    for match in _SPEAKER_TAG_RE.finditer(raw):
        gender, accent, age, register = match.groups()
        tag = match.group(0).strip()
        if tag in seen:
            continue
        seen[tag] = {
            "tag": tag,
            "gender": gender.upper() if gender else None,
            "accent": accent if accent else None,
            "age": age.lower() if age else None,
            "register": register.lower() if register else None,
        }
    return list(seen.values())


# ── Metadata table ──────────────────────────────────────────────────────────


# Header-text → metadata-field map. Lookup is case-insensitive +
# whitespace-trimmed. Andy may reorder rows; this map is resilient.
_METADATA_HEADERS: dict[str, str] = {
    "test id":          "test_id",
    "test identifier":  "test_id",
    "version":          "version",
    "band target":      "band_target",
    "target band":      "band_target",
    "themes":           "themes_raw",
    "theme":            "themes_raw",
    "accent profile":   "accent_profile_raw",
    "accents":          "accent_profile_raw",
    "total words":      "total_words",
    "word count":       "total_words",
    "source":           "source_format",
    "source format":    "source_format",
    "created at":       "created_at_source",
    "date":             "created_at_source",
}


def parse_metadata_table(tables: list[list[list[str]]]) -> dict[str, Any]:
    """Find the metadata table (any 2-column table whose first column
    contains at least 3 known headers) and extract its fields.

    Returns a dict with normalized keys. ``themes`` is parsed into
    ``{s1, s2, s3, s4}``; ``accent_profile`` into ``list[str]``;
    ``band_target`` to ``float``; ``total_words`` to ``int``.
    """
    raw: dict[str, str] = {}
    for table in tables:
        hits = 0
        candidate: dict[str, str] = {}
        for row in table:
            if len(row) < 2:
                continue
            key = row[0].strip().lower()
            val = row[1].strip()
            if key in _METADATA_HEADERS and val:
                candidate[_METADATA_HEADERS[key]] = val
                hits += 1
        if hits >= 3:                           # require ≥3 known fields
            raw = candidate
            break

    out: dict[str, Any] = {
        "test_id":           raw.get("test_id"),
        "version":           raw.get("version") or "1.0",
        "band_target":       None,
        "themes":            {},
        "accent_profile":    [],
        "total_words":       None,
        "source_format":     raw.get("source_format") or "cambridge_ielts_docx",
        "created_at_source": raw.get("created_at_source"),
    }

    if "band_target" in raw:
        m = re.search(r"(\d+(?:\.\d+)?)", raw["band_target"])
        if m:
            try:
                out["band_target"] = float(m.group(1))
            except ValueError:
                pass

    if "total_words" in raw:
        m = re.search(r"(\d[\d,]*)", raw["total_words"])
        if m:
            try:
                out["total_words"] = int(m.group(1).replace(",", ""))
            except ValueError:
                pass

    if "themes_raw" in raw:
        out["themes"] = _parse_themes(raw["themes_raw"])

    if "accent_profile_raw" in raw:
        out["accent_profile"] = _parse_accent_profile(raw["accent_profile_raw"])

    return out


def _parse_themes(raw: str) -> dict[str, str]:
    """Parse 'Section 1: X; Section 2: Y; Section 3: Z; Section 4: W' or
    newline-separated variants into a ``{s1, s2, s3, s4}`` dict.
    """
    out: dict[str, str] = {}
    parts = re.split(r"[;\n]", raw)
    for part in parts:
        m = re.match(r"\s*Section\s*([1-4])\s*[:\-—]\s*(.+?)\s*$", part, re.IGNORECASE)
        if m:
            out[f"s{m.group(1)}"] = m.group(2).strip()
    return out


def _parse_accent_profile(raw: str) -> list[str]:
    """Parse 'BrE, AusE' / 'BrE; AmE' / 'BrE / AusE' → list[str]."""
    parts = re.split(r"[,;/|]", raw)
    return [p.strip() for p in parts if p.strip()]


# ── Section split ───────────────────────────────────────────────────────────


def split_sections(script_text: str) -> dict[int, str]:
    """Split a Script+AnswerKey transcript into per-section raw text.

    Returns ``{section_num: section_raw_text}`` for sections found.
    Empty/missing sections are simply absent from the dict.
    """
    lines = script_text.splitlines()
    headers: list[tuple[int, int]] = []          # (line_idx, section_num)
    for i, line in enumerate(lines):
        m = _SECTION_HEADER_RE.match(line)
        if m:
            headers.append((i, int(m.group(1))))

    out: dict[int, str] = {}
    for idx, (line_idx, section_num) in enumerate(headers):
        end_idx = headers[idx + 1][0] if idx + 1 < len(headers) else len(lines)
        body = "\n".join(lines[line_idx + 1:end_idx]).strip()
        if body:
            out[section_num] = body
    return out


# ── Question Paper parse ────────────────────────────────────────────────────


def _classify_instruction(instruction: str) -> str:
    """Map a Question Paper instruction line to a question-type slug."""
    for pattern, slug in _INSTRUCTION_HINTS:
        if pattern.search(instruction):
            return slug
    return "unknown"


def parse_question_paper(qp_text: str) -> dict[int, list[dict[str, Any]]]:
    """Parse the Question Paper text into a per-section list of questions.

    Returns ``{section_num: [{q_num, prompt, q_type, options?}, ...]}``.

    Recognises:
      * "Section N" headers (case-insensitive)
      * One or more instruction lines per section ("Complete the form ..."
        / "Choose the correct letter, A, B or C") — last instruction
        before a question wins
      * "Daniel Brennan (Example)" or any line containing "(Example)" is
        skipped (Cambridge worked-example convention)
      * Numbered question lines "1 ……" / "1. ……" / "1\t……"
      * Sub-bullet MCQ options "A …" / "B …" / "C …" / "D …" / ... → up
        to 8 letter options for plan/map labelling
    """
    out: dict[int, list[dict[str, Any]]] = {1: [], 2: [], 3: [], 4: []}
    section_num = 0
    instruction: str = ""
    current_q_type = "unknown"
    pending_options: list[tuple[str, str]] = []
    last_question: dict[str, Any] | None = None

    section_header_re = re.compile(
        r"^\s*Section\s+([1-4])\s*$", re.IGNORECASE,
    )
    option_line_re = re.compile(
        r"^\s*([A-H])\s*[\.\):]\s+(.+?)\s*$",
    )

    def _flush_options() -> None:
        if last_question is not None and pending_options:
            last_question.setdefault("options", []).extend(
                {"letter": L, "text": T} for L, T in pending_options
            )
        pending_options.clear()

    for raw_line in qp_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if "(example)" in line.lower():         # skip Cambridge worked example
            continue

        m_sec = section_header_re.match(line)
        if m_sec:
            _flush_options()
            section_num = int(m_sec.group(1))
            last_question = None
            instruction = ""
            current_q_type = "unknown"
            continue

        # Instruction lines (no leading digit, no leading letter+bullet)
        m_q = _QUESTION_NUMBER_LINE_RE.match(line)
        m_opt = option_line_re.match(line)

        if m_q and section_num:
            _flush_options()
            q_num = int(m_q.group(1))
            prompt = m_q.group(2).strip()
            entry: dict[str, Any] = {
                "q_num": q_num,
                "prompt": prompt,
                "q_type": current_q_type,
            }
            out[section_num].append(entry)
            last_question = entry
            continue

        if m_opt and section_num and last_question is not None:
            pending_options.append((m_opt.group(1), m_opt.group(2).strip()))
            continue

        # Treat anything else as a potential instruction line.
        candidate_type = _classify_instruction(line)
        if candidate_type != "unknown":
            _flush_options()
            instruction = line
            current_q_type = candidate_type

    _flush_options()
    return out


# ── Answer Key parse ────────────────────────────────────────────────────────


def parse_answer_key(tables: list[list[list[str]]]) -> dict[int, list[dict[str, Any]]]:
    """Parse all answer-key tables in the Script+AnswerKey doc.

    Heuristic: a table is an answer key if its header row contains
    ``question`` and ``answer`` columns. Trap-mechanism column is optional
    (third column, anything goes). Multiple answer-key tables (one per
    section) are concatenated then grouped by question number range.
    """
    rows_collected: list[tuple[int, str, str]] = []  # (q_num, answer, trap)

    for table in tables:
        if not table or len(table) < 2:
            continue
        header = [c.strip().lower() for c in table[0]]
        if not any("question" in h or h == "q" or h == "q#" for h in header):
            continue
        if not any("answer" in h for h in header):
            continue
        # Locate column indices (defensive — Andy may reorder columns).
        try:
            q_col = next(
                i for i, h in enumerate(header)
                if "question" in h or h in ("q", "q#")
            )
            a_col = next(i for i, h in enumerate(header) if "answer" in h)
        except StopIteration:                   # pragma: no cover
            continue
        trap_col: int | None = None
        for i, h in enumerate(header):
            if "trap" in h or "mechanism" in h:
                trap_col = i
                break

        for row in table[1:]:
            if len(row) <= max(q_col, a_col):
                continue
            q_cell = row[q_col].strip()
            a_cell = row[a_col].strip()
            if not q_cell or not a_cell:
                continue
            if "(example)" in a_cell.lower():
                continue
            m = re.match(r"^(\d{1,2})", q_cell)
            if not m:
                continue
            q_num = int(m.group(1))
            trap = (
                row[trap_col].strip()
                if trap_col is not None and trap_col < len(row)
                else ""
            )
            rows_collected.append((q_num, a_cell, trap))

    # Group by Cambridge section convention: Q1-10 → s1, Q11-20 → s2,
    # Q21-30 → s3, Q31-40 → s4.
    out: dict[int, list[dict[str, Any]]] = {1: [], 2: [], 3: [], 4: []}
    for q_num, answer, trap in rows_collected:
        section_num = ((q_num - 1) // 10) + 1
        if 1 <= section_num <= 4:
            out[section_num].append({
                "q_num":           q_num,
                "answer":          answer,
                "trap_mechanisms": [t.strip() for t in re.split(r"[;,]", trap) if t.strip()],
            })
    for section_num in out:
        out[section_num].sort(key=lambda r: r["q_num"])
    return out


# ── Exercise grouping ───────────────────────────────────────────────────────


_Q_TYPE_TO_EXERCISE: dict[str, str] = {
    "dictation_gap_fill":     "dictation",
    "dictation_short_answer": "dictation",
    "mcq_3option":            "mcq",
    "mcq_letter_label":       "mcq",
    "unknown":                "dictation",      # safe default — admin can edit
}


def build_exercises(
    questions: list[dict[str, Any]],
    answers: list[dict[str, Any]],
    section_num: int,
) -> list[dict[str, Any]]:
    """Group consecutive same-type questions into exercise payloads.

    One exercise row per contiguous run of identical ``q_type`` values.
    Each payload carries ``{questions: [...], answers: [{q_num, answer,
    trap_mechanisms}], variant}`` — admin reviews + edits in Sprint 13.4
    UI; student delivery comes from Sprint 13.5 student-test renderer.
    """
    answer_by_q = {a["q_num"]: a for a in answers}
    exercises: list[dict[str, Any]] = []
    if not questions:
        return exercises

    run_type = questions[0]["q_type"]
    run: list[dict[str, Any]] = []

    def _flush() -> None:
        if not run:
            return
        exercise_type = _Q_TYPE_TO_EXERCISE.get(run_type, "dictation")
        payload_questions = [
            {
                "q_num":   q["q_num"],
                "prompt":  q["prompt"],
                **({"options": q["options"]} if q.get("options") else {}),
            }
            for q in run
        ]
        payload_answers = [
            answer_by_q[q["q_num"]]
            for q in run
            if q["q_num"] in answer_by_q
        ]
        exercises.append({
            "exercise_type": exercise_type,
            "variant":       run_type,
            "section_num":   section_num,
            "order_num":     len(exercises) + 1,
            "payload": {
                "variant":   run_type,
                "questions": payload_questions,
                "answers":   payload_answers,
            },
        })

    for q in questions:
        if q["q_type"] != run_type:
            _flush()
            run = [q]
            run_type = q["q_type"]
        else:
            run.append(q)
    _flush()
    return exercises


# ── Accent + CEFR inference ─────────────────────────────────────────────────


_ACCENT_TO_TAG = {
    "bre":  "uk_rp",
    "amer": "us_general",
    "ame":  "us_general",
    "use":  "us_general",
    "us":   "us_general",
    "ause": "au",
    "aue":  "au",
    "au":   "au",
    "cae":  "ca",
    "ca":   "ca",
}


def infer_accent_tag(speakers: list[dict[str, str | None]]) -> str:
    """Map the section's speaker mix to one of the legal accent_tag
    enum values: us_general | uk_rp | au | ca | other. Mixed accents
    default to 'other'.
    """
    accents = {
        (sp.get("accent") or "").lower()
        for sp in speakers
        if sp.get("accent")
    }
    if not accents:
        return "other"
    mapped = {
        _ACCENT_TO_TAG.get(a, "other")
        for a in accents
    }
    if len(mapped) == 1:
        return mapped.pop()
    return "other"


# Cambridge band → CEFR coarse map. Half-bands round down to coarser.
def infer_cefr_level(band_target: float | None) -> str | None:
    """Map a Cambridge band score (e.g. 5.5) to CEFR (e.g. B1)."""
    if band_target is None:
        return None
    if band_target >= 8.5:
        return "C2"
    if band_target >= 7.0:
        return "C1"
    if band_target >= 5.5:
        return "B2"
    if band_target >= 4.0:
        return "B1"
    return "A2"


# ── Public entry point ──────────────────────────────────────────────────────


def parse_listening_test(
    question_paper_bytes: bytes,
    script_answerkey_bytes: bytes,
) -> dict[str, Any]:
    """Parse Andy's 2-file Cambridge IELTS DOCX bundle into structured
    data ready for ``listening_tests`` + ``listening_content`` ingest.

    Returns:
        {
          "test_metadata": {...},
          "sections": [
            {
              "section_num": 1,
              "title": "Cookery class enrolment",
              "transcript_raw": "...",
              "transcript_clean": "...",
              "speakers": [...],
              "word_count": int,
              "accent_tag": "uk_rp",
              "cefr_level": "B1",
              "ielts_section": 1,
              "questions": [...],
              "answers": [...],
              "exercises": [...],
            }, ...
          ],
          "warnings": [str, ...],
          "errors":   [str, ...],
        }
    """
    qp_text, _qp_tables = _extract_docx(question_paper_bytes)
    sa_text, sa_tables = _extract_docx(script_answerkey_bytes)
    return parse_from_text(qp_text, sa_text, sa_tables)


def parse_from_text(
    qp_text: str,
    sa_text: str,
    sa_tables: list[list[list[str]]],
) -> dict[str, Any]:
    """Pure-text entry point — used by tests that synthesize input
    without going through python-docx.
    """
    warnings: list[str] = []
    errors: list[str] = []

    metadata = parse_metadata_table(sa_tables)
    if not metadata.get("test_id"):
        errors.append(
            "Không tìm thấy Test ID trong metadata table. Kiểm tra hàng "
            "đầu tiên của bảng metadata."
        )

    section_texts = split_sections(sa_text)
    if len(section_texts) != 4:
        warnings.append(
            f"Phát hiện {len(section_texts)} sections trong transcript "
            f"(kỳ vọng 4). Section thiếu sẽ bị bỏ qua."
        )

    qp_questions = parse_question_paper(qp_text)
    answer_key = parse_answer_key(sa_tables)

    sections_out: list[dict[str, Any]] = []
    total_word_count = 0

    for section_num in (1, 2, 3, 4):
        raw = section_texts.get(section_num, "")
        if not raw:
            warnings.append(f"Section {section_num} không có transcript trong file Script.")
            continue

        clean = strip_markers(raw)
        speakers = parse_speakers(raw)
        word_count = len(re.findall(r"\b\w+\b", clean))
        total_word_count += word_count

        questions = qp_questions.get(section_num, [])
        answers = answer_key.get(section_num, [])

        if not questions:
            warnings.append(f"Section {section_num} không có questions trong file Question Paper.")
        if len(answers) != len(questions):
            warnings.append(
                f"Section {section_num} mismatch: {len(questions)} questions "
                f"vs {len(answers)} answers."
            )

        exercises = build_exercises(questions, answers, section_num)

        themes = metadata.get("themes") or {}
        theme_text = themes.get(f"s{section_num}", "")

        accent_tag = infer_accent_tag(speakers)
        cefr_level = infer_cefr_level(metadata.get("band_target"))

        sections_out.append({
            "section_num":      section_num,
            "title":            theme_text or f"Section {section_num}",
            "theme":            theme_text,
            "transcript_raw":   raw,
            "transcript_clean": clean,
            "speakers":         speakers,
            "word_count":       word_count,
            "accent_tag":       accent_tag,
            "cefr_level":       cefr_level,
            "ielts_section":    section_num,
            "questions":        questions,
            "answers":          answers,
            "exercises":        exercises,
        })

    # Word-count sanity: warn if parser total drifts >10% from metadata.
    if metadata.get("total_words") and total_word_count:
        declared = float(metadata["total_words"])
        drift = abs(total_word_count - declared) / declared
        if drift > 0.10:
            warnings.append(
                f"Word count drift {int(drift * 100)}%: parser counted "
                f"{total_word_count}, metadata declared {int(declared)}."
            )

    return {
        "test_metadata": metadata,
        "sections":      sections_out,
        "warnings":      warnings,
        "errors":        errors,
        "_total_word_count": total_word_count,
    }


# ── Convenience: section → listening_content row payload ────────────────────


def section_to_content_payload(
    section: dict[str, Any],
    test_id_uuid: str,
    test_metadata: dict[str, Any],
) -> dict[str, Any]:
    """Build the ``listening_content`` INSERT payload for one section.

    Audio fields use the Sprint 13.3.1 placeholder pattern
    (audio_storage_path=NULL, duration=0, size=0, status='draft').
    Andy uploads MP3s later via the existing bulk-upload flow
    (Sprint 13.5 wires a "link audio to section" UI).
    """
    test_external = test_metadata.get("test_id") or "test"
    theme_slug = re.sub(r"[^a-z0-9]+", "-", (section.get("theme") or "").lower()).strip("-")
    topic_tags = [test_external, f"section-{section['section_num']}"]
    if theme_slug:
        topic_tags.append(theme_slug)

    return {
        "source_type":            "test_section",
        "test_id":                test_id_uuid,
        "section_num":            section["section_num"],
        "audio_storage_path":     None,
        "audio_duration_seconds": 0,
        "audio_size_bytes":       0,
        "title":                  f"{test_external} — Section {section['section_num']}: "
                                  f"{section.get('theme') or 'Untitled'}",
        "transcript":             section["transcript_clean"],
        "accent_tag":             section["accent_tag"],
        "cefr_level":             section.get("cefr_level"),
        "ielts_section":          section["ielts_section"],
        "topic_tags":             topic_tags,
        "status":                 "draft",
        "is_premium":             False,
        "metadata": {
            "source_format":   test_metadata.get("source_format") or "cambridge_ielts_docx",
            "speakers":        section["speakers"],
            "raw_transcript":  section["transcript_raw"],
            "word_count":      section["word_count"],
            "theme":           section.get("theme"),
        },
    }
